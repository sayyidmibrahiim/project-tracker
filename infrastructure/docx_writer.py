"""Tiptap JSON -> .docx renderer for the RTE document pipeline (D-0012).

Renders a document `source` dict (see services/rte_document_service.py for the
source.json schema) into a real Word file with python-docx. Writing is atomic:
render to a sibling tmp file, validate it re-opens, then os.replace() onto the
target. A locked target (Word has the file open) raises DocxTargetLockedError
and leaves the previous .docx untouched.

Mapping fidelity follows _docs/flow-tiptap.md §12 "Jalur B": every toolbar
feature has an explicit mapping; unmapped nodes degrade to plain paragraphs,
never crash the export.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Any, Callable

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image as _DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


class DocxTargetLockedError(RuntimeError):
    """Raised when the target .docx cannot be replaced (open in Word)."""


DEFAULT_DOCUMENT_SETTINGS: dict[str, Any] = {
    "page_format": "A4",
    "margin_top_mm": 12.7,
    "margin_right_mm": 12.7,
    "margin_bottom_mm": 12.7,
    "margin_left_mm": 12.7,
    "default_font_family": "Times New Roman",
    # 13.5pt = 18px editor default × 0.75 (editor↔Word WYSIWYG).
    "default_font_size_pt": 13.5,
    "line_height": 1.15,
}


def content_width_mm(settings: dict[str, Any]) -> float:
    return max(
        210.0
        - float(settings.get("margin_left_mm", 12.7))
        - float(settings.get("margin_right_mm", 12.7)),
        40.0,
    )


_ALIGN = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

_LIST_STYLES = {"bulletList": "List Bullet", "orderedList": "List Number"}


def _hex_to_rgb(value: Any) -> RGBColor | None:
    if not isinstance(value, str):
        return None
    v = value.strip().lstrip("#")
    if len(v) == 3:
        v = "".join(ch * 2 for ch in v)
    if len(v) != 6:
        return None
    try:
        return RGBColor.from_string(v.upper())
    except (ValueError, TypeError):
        return None


def _hex_fill(value: Any) -> str | None:
    rgb = _hex_to_rgb(value)
    return str(rgb) if rgb is not None else None


def _px_to_pt(px: Any) -> float | None:
    """Editor font sizes are CSS px strings like '14px'; Word wants points."""
    try:
        if isinstance(px, str):
            px = px.strip().removesuffix("px")
        return float(px) * 0.75
    except (ValueError, TypeError):
        return None


def _shade(element: Any, fill: str) -> None:
    """Apply a w:shd fill to a run rPr / paragraph pPr / cell tcPr element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill)
    element.append(shd)


def _first_family(value: Any) -> str | None:
    if not isinstance(value, str) or not value.strip():
        return None
    return value.split(",")[0].strip().strip("'\"")


class _Renderer:
    def __init__(self, document: Document, assets_dir: Path, settings: dict[str, Any]) -> None:
        self.doc = document
        self.assets_dir = assets_dir
        self.settings = settings
        self.content_width_px = content_width_mm(settings) / 25.4 * 96.0
        self._renderers: dict[str, Callable[[dict, Any], None]] = {
            "paragraph": self._paragraph,
            "heading": self._heading,
            "bulletList": self._list,
            "orderedList": self._list,
            "taskList": self._task_list,
            "blockquote": self._blockquote,
            "codeBlock": self._code_block,
            "horizontalRule": self._horizontal_rule,
            "table": self._table,
            "image": self._block_image,
        }

    # ── Block dispatch ──

    def render_body(self, nodes: list[dict]) -> None:
        for node in nodes or []:
            self._render_block(node)

    def _render_block(self, node: dict, *, container: Any = None) -> None:
        handler = self._renderers.get(node.get("type", ""))
        if handler is not None:
            handler(node, container)
        else:
            # Unknown block: degrade to a paragraph with its inline text.
            self._paragraph(node, container)

    def _add_paragraph(self, container: Any, style: str | None = None):
        target = container if container is not None else self.doc
        p = target.add_paragraph()
        if style:
            try:
                p.style = style
            except KeyError:
                pass
        return p

    # ── Blocks ──

    def _paragraph(self, node: dict, container: Any, style: str | None = None) -> None:
        p = self._add_paragraph(container, style)
        self._apply_para_attrs(p, node.get("attrs") or {})
        p.paragraph_format.space_after = Pt(8)
        line = self.settings.get("line_height")
        if isinstance(line, (int, float)) and line > 0:
            p.paragraph_format.line_spacing = float(line)
        self._render_inlines(p, node.get("content") or [])

    def _heading(self, node: dict, container: Any) -> None:
        attrs = node.get("attrs") or {}
        level = attrs.get("level")
        level = level if isinstance(level, int) and 1 <= level <= 6 else 1
        p = self._add_paragraph(container, f"Heading {level}")
        self._apply_para_attrs(p, attrs)
        self._render_inlines(p, node.get("content") or [])

    def _list(self, node: dict, container: Any, depth: int = 0) -> None:
        base = _LIST_STYLES.get(node.get("type", "bulletList"), "List Bullet")
        style = base if depth == 0 else f"{base} {min(depth + 1, 3)}"
        for item in node.get("content") or []:
            for child in item.get("content") or []:
                ctype = child.get("type")
                if ctype in _LIST_STYLES:
                    self._list(child, container, depth + 1)
                elif ctype == "paragraph":
                    p = self._add_paragraph(container, style)
                    self._render_inlines(p, child.get("content") or [])
                else:
                    self._render_block(child, container=container)

    def _task_list(self, node: dict, container: Any) -> None:
        for item in node.get("content") or []:
            checked = bool((item.get("attrs") or {}).get("checked"))
            p = self._add_paragraph(container, "List Bullet")
            p.add_run("☑ " if checked else "☐ ")
            for child in item.get("content") or []:
                if child.get("type") == "paragraph":
                    self._render_inlines(p, child.get("content") or [])

    def _blockquote(self, node: dict, container: Any) -> None:
        for child in node.get("content") or []:
            p = self._add_paragraph(container)
            p.paragraph_format.left_indent = Mm(10)
            self._render_inlines(p, child.get("content") or [], force_italic=True)

    def _code_block(self, node: dict, container: Any) -> None:
        p = self._add_paragraph(container)
        _shade(p._p.get_or_add_pPr(), "F3F3F3")
        text = "".join(
            child.get("text", "") for child in node.get("content") or [] if child.get("type") == "text"
        )
        run = p.add_run(text)
        run.font.name = "Consolas"
        run.font.size = Pt(10)

    def _horizontal_rule(self, node: dict, container: Any) -> None:
        p = self._add_paragraph(container)
        p_pr = p._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "auto")
        borders.append(bottom)
        p_pr.append(borders)

    def _block_image(self, node: dict, container: Any) -> None:
        p = self._add_paragraph(container)
        self._image_run(p, node)

    # ── Table ──

    def _table(self, node: dict, container: Any) -> None:
        rows = [r for r in node.get("content") or [] if r.get("type") == "tableRow"]
        if not rows:
            return
        n_cols = 0
        for cell in rows[0].get("content") or []:
            n_cols += int((cell.get("attrs") or {}).get("colspan") or 1)
        n_cols = max(n_cols, 1)
        target = container if container is not None else self.doc
        table = target.add_table(rows=len(rows), cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False
        try:
            table.style = "Table Grid"
        except KeyError:
            pass

        col_widths_px: list[float | None] = [None] * n_cols
        occupied: set[tuple[int, int]] = set()

        for r, row in enumerate(rows):
            c = 0
            for cell_node in row.get("content") or []:
                while (r, c) in occupied and c < n_cols:
                    c += 1
                if c >= n_cols:
                    break
                attrs = cell_node.get("attrs") or {}
                colspan = int(attrs.get("colspan") or 1)
                rowspan = int(attrs.get("rowspan") or 1)
                colwidth = attrs.get("colwidth")
                if isinstance(colwidth, list):
                    for i, w in enumerate(colwidth[:colspan]):
                        if isinstance(w, (int, float)) and c + i < n_cols:
                            col_widths_px[c + i] = float(w)

                cell = table.cell(r, c)
                if colspan > 1 and c + colspan - 1 < n_cols:
                    cell = cell.merge(table.cell(r, c + colspan - 1))
                if rowspan > 1 and r + rowspan - 1 < len(rows):
                    cell = cell.merge(table.cell(r + rowspan - 1, c))
                for rr in range(r, min(r + rowspan, len(rows))):
                    for cc in range(c, min(c + colspan, n_cols)):
                        occupied.add((rr, cc))

                is_header = cell_node.get("type") == "tableHeader"
                if is_header:
                    _shade(cell._tc.get_or_add_tcPr(), "E7E6E6")

                # Cells come with one empty default paragraph; fill it first.
                first = True
                for child in cell_node.get("content") or []:
                    if child.get("type") == "paragraph" and first:
                        p = cell.paragraphs[0]
                        self._apply_para_attrs(p, child.get("attrs") or {})
                        self._render_inlines(p, child.get("content") or [], force_bold=is_header)
                        first = False
                    elif child.get("type") == "paragraph":
                        p = cell.add_paragraph()
                        self._apply_para_attrs(p, child.get("attrs") or {})
                        self._render_inlines(p, child.get("content") or [], force_bold=is_header)
                    else:
                        self._render_block(child, container=cell)
                c += colspan

        specified_width_px = sum(px for px in col_widths_px if px is not None)
        if specified_width_px > self.content_width_px:
            scale = self.content_width_px / specified_width_px
            col_widths_px = [px * scale if px is not None else None for px in col_widths_px]

        # Column widths: px -> dxa (1 px ≈ 15 dxa at 96dpi/20 twips-per-pt).
        for idx, px in enumerate(col_widths_px):
            if px is None:
                continue
            dxa = str(int(px * 15))
            for row_cells in table.rows:
                try:
                    tc = row_cells.cells[idx]._tc
                except IndexError:
                    continue
                tc_pr = tc.get_or_add_tcPr()
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is None:
                    tc_w = OxmlElement("w:tcW")
                    tc_pr.append(tc_w)
                tc_w.set(qn("w:w"), dxa)
                tc_w.set(qn("w:type"), "dxa")

    # ── Inlines ──

    def _render_inlines(
        self,
        paragraph: Any,
        nodes: list[dict],
        *,
        force_bold: bool = False,
        force_italic: bool = False,
    ) -> None:
        for node in nodes:
            ntype = node.get("type")
            if ntype == "text":
                marks = node.get("marks") or []
                link = next((m for m in marks if m.get("type") == "link"), None)
                if link is not None:
                    self._hyperlink(paragraph, node, link, force_bold, force_italic)
                else:
                    run = paragraph.add_run(node.get("text", ""))
                    self._apply_marks(run, marks, force_bold, force_italic)
            elif ntype == "hardBreak":
                paragraph.add_run().add_break()
            elif ntype == "image":
                self._image_run(paragraph, node)
            elif node.get("content"):
                self._render_inlines(
                    paragraph, node["content"], force_bold=force_bold, force_italic=force_italic
                )

    def _apply_marks(
        self, run: Any, marks: list[dict], force_bold: bool = False, force_italic: bool = False
    ) -> None:
        font = run.font
        font.name = _first_family(self.settings.get("default_font_family")) or "Times New Roman"
        size_pt = self.settings.get("default_font_size_pt")
        if isinstance(size_pt, (int, float)):
            font.size = Pt(float(size_pt))
        if force_bold:
            run.bold = True
        if force_italic:
            run.italic = True
        for mark in marks:
            mtype = mark.get("type")
            attrs = mark.get("attrs") or {}
            if mtype == "bold":
                run.bold = True
            elif mtype == "italic":
                run.italic = True
            elif mtype == "underline":
                run.underline = True
            elif mtype == "strike":
                font.strike = True
            elif mtype == "subscript":
                font.subscript = True
            elif mtype == "superscript":
                font.superscript = True
            elif mtype == "code":
                font.name = "Consolas"
            elif mtype == "highlight":
                fill = _hex_fill(attrs.get("color")) or "FFFF00"
                _shade(run._r.get_or_add_rPr(), fill)
            elif mtype == "textStyle":
                rgb = _hex_to_rgb(attrs.get("color"))
                if rgb is not None:
                    font.color.rgb = rgb
                family = _first_family(attrs.get("fontFamily"))
                if family:
                    font.name = family
                pt = _px_to_pt(attrs.get("fontSize"))
                if pt:
                    font.size = Pt(pt)
        # rFonts ascii/hAnsi/cs must all be set or Word falls back per-script.
        if font.name:
            rpr = run._r.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            for attr in ("w:ascii", "w:hAnsi", "w:cs"):
                rfonts.set(qn(attr), font.name)

    def _hyperlink(
        self, paragraph: Any, node: dict, link_mark: dict, force_bold: bool, force_italic: bool
    ) -> None:
        href = (link_mark.get("attrs") or {}).get("href") or ""
        try:
            from docx.opc.constants import RELATIONSHIP_TYPE as RT

            r_id = paragraph.part.relate_to(href, RT.HYPERLINK, is_external=True)
        except (ValueError, KeyError):
            run = paragraph.add_run(node.get("text", ""))
            self._apply_marks(run, node.get("marks") or [], force_bold, force_italic)
            return
        hyperlink = OxmlElement("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)
        run = paragraph.add_run(node.get("text", ""))
        self._apply_marks(
            run,
            [m for m in node.get("marks") or [] if m.get("type") != "link"],
            force_bold,
            force_italic,
        )
        run.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
        run.underline = True
        run._r.getparent().remove(run._r)
        hyperlink.append(run._r)
        paragraph._p.append(hyperlink)

    def _image_run(self, paragraph: Any, node: dict) -> None:
        attrs = node.get("attrs") or {}
        asset_id = attrs.get("assetId")
        src = attrs.get("src") or ""
        file_name = None
        if isinstance(src, str) and src.startswith("asset://"):
            file_name = src.removeprefix("asset://")
        elif asset_id:
            file_name = next(
                (p.name for p in sorted(self.assets_dir.glob(f"{asset_id}.*"))), None
            ) if self.assets_dir.is_dir() else None
        path = (self.assets_dir / file_name) if file_name else None
        if path is None or not path.is_file():
            run = paragraph.add_run(f"[missing image {asset_id or src}]")
            run.italic = True
            return
        width = attrs.get("width")
        kwargs = {}
        max_width = Inches(self.content_width_px / 96.0)
        if isinstance(width, (int, float)) and width > 0:
            kwargs["width"] = min(Inches(float(width) / 96.0), max_width)
        else:
            try:
                probe = _DocxImage.from_file(str(path))
                natural = Inches(probe.px_width / float(probe.horz_dpi or 96))
                if natural > max_width:
                    kwargs["width"] = max_width
            except Exception:
                pass
        try:
            paragraph.add_run().add_picture(str(path), **kwargs)
        except Exception:
            run = paragraph.add_run(f"[unreadable image {path.name}]")
            run.italic = True

    # ── Page/paragraph attrs ──

    def _apply_para_attrs(self, paragraph: Any, attrs: dict) -> None:
        align = _ALIGN.get(str(attrs.get("textAlign") or "").lower())
        if align is not None:
            paragraph.alignment = align


def _apply_page_setup(document: Document, settings: dict[str, Any]) -> None:
    section = document.sections[0]
    if str(settings.get("page_format", "A4")).upper() == "A4":
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    section.top_margin = Mm(float(settings.get("margin_top_mm", 12.7)))
    section.right_margin = Mm(float(settings.get("margin_right_mm", 12.7)))
    section.bottom_margin = Mm(float(settings.get("margin_bottom_mm", 12.7)))
    section.left_margin = Mm(float(settings.get("margin_left_mm", 12.7)))
    style = document.styles["Normal"]
    family = _first_family(settings.get("default_font_family"))
    if family:
        style.font.name = family
    size_pt = settings.get("default_font_size_pt")
    if isinstance(size_pt, (int, float)):
        style.font.size = Pt(float(size_pt))


def export_source_to_docx(source: dict, assets_dir: Path, target_docx: Path) -> None:
    """Render source['content'] into target_docx atomically.

    Raises DocxTargetLockedError when the final os.replace fails because the
    target is open in another program (Word lock). Any other failure removes
    the tmp file and re-raises; the previous target file is never touched.
    """
    settings = {**DEFAULT_DOCUMENT_SETTINGS, **(source.get("document_settings") or {})}
    content = source.get("content") or {}

    document = Document()
    _apply_page_setup(document, settings)
    _Renderer(document, assets_dir, settings).render_body(content.get("content") or [])

    tmp = target_docx.with_name(f".{target_docx.name}.tmp")
    try:
        document.save(str(tmp))
        Document(str(tmp))  # validation: the rendered file must re-open
        os.replace(tmp, target_docx)
    except PermissionError as exc:
        tmp.unlink(missing_ok=True)
        raise DocxTargetLockedError(str(target_docx)) from exc
    except Exception:
        tmp.unlink(missing_ok=True)
        raise
