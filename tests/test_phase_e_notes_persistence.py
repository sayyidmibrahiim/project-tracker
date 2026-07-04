"""Phase E.1 — Notes persistence via notes.md file (PRD-correct).

Tests prove:
- notes_get reads from notes.md file, not from project_data.json
- notes_update writes to notes.md file and persists
- missing notes.md returns empty string
- update creates notes.md if absent
- unrelated mutations remain unavailable
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import pytest

from core.enums import CRState
from core.models import ProjectMetadata
from infrastructure.metadata_store import MetadataStore
from infrastructure.settings_store import SettingsStore


@pytest.fixture
def temp_project():
    """Temp project with notes.md file."""
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        year_dir = root / "2024"
        state_dir = year_dir / "UAT_PREPARE"
        proj_dir = state_dir / "test-project"
        proj_dir.mkdir(parents=True)

        metadata = ProjectMetadata(
            project_name="test-project",
            cr_link="https://cr.example.com/CR-999",
            cr_state=CRState.APPROVED,
        )
        store = MetadataStore()
        store.write(proj_dir, metadata)

        # Create notes.md with content
        (proj_dir / "notes.md").write_text("existing notes content")

        settings = SettingsStore(config_dir=root / "config")
        from dataclasses import replace

        updated = replace(settings.read(), root_folder=root)
        settings.write(updated)

        yield {
            "root": root,
            "project_path": proj_dir,
            "settings_store": settings,
        }


@pytest.fixture
def temp_project_no_notes():
    """Temp project WITHOUT notes.md file."""
    with tempfile.TemporaryDirectory() as tmp:
        root = Path(tmp)
        year_dir = root / "2024"
        state_dir = year_dir / "UAT_PREPARE"
        proj_dir = state_dir / "test-project-no-notes"
        proj_dir.mkdir(parents=True)

        metadata = ProjectMetadata(project_name="test-project-no-notes")
        store = MetadataStore()
        store.write(proj_dir, metadata)

        settings = SettingsStore(config_dir=root / "config")
        from dataclasses import replace

        updated = replace(settings.read(), root_folder=root)
        settings.write(updated)

        yield {
            "root": root,
            "project_path": proj_dir,
            "settings_store": settings,
        }


@pytest.fixture
def js_api(temp_project):
    """JsApi from create_js_api with temp settings."""
    from project_tracker import app_web

    return app_web.create_js_api(settings_store=temp_project["settings_store"])


@pytest.fixture
def js_api_no_notes(temp_project_no_notes):
    """JsApi from create_js_api for project without notes.md."""
    from project_tracker import app_web

    return app_web.create_js_api(settings_store=temp_project_no_notes["settings_store"])


# ── notes_get reads from notes.md file ──────────────────────────────────


def test_notes_get_reads_from_notes_md_file(js_api, temp_project):
    """notes_get must read content from notes.md, not from project_data.json."""
    path = str(temp_project["project_path"])
    result = js_api.notes_get(path)
    assert result["ok"] is True
    assert result["data"] == "existing notes content"


def test_notes_get_returns_empty_when_no_notes_md(js_api_no_notes, temp_project_no_notes):
    """notes_get returns empty string when notes.md does not exist."""
    path = str(temp_project_no_notes["project_path"])
    result = js_api_no_notes.notes_get(path)
    assert result["ok"] is True
    assert result["data"] == ""


# ── notes_update writes to notes.md file ─────────────────────────────────


def test_notes_update_no_longer_service_unavailable(js_api, temp_project):
    """notes_update must not return SERVICE_UNAVAILABLE."""
    path = str(temp_project["project_path"])
    result = js_api.notes_update(path, "new notes")
    assert result["ok"] is True, f"Expected ok=True, got {result}"


def test_notes_update_persists_to_notes_md(js_api, temp_project):
    """notes_update writes content to notes.md file."""
    path = str(temp_project["project_path"])
    js_api.notes_update(path, "updated content")

    notes_file = temp_project["project_path"] / "notes.md"
    assert notes_file.read_text() == "updated content"


def test_notes_update_creates_notes_md_if_absent(js_api_no_notes, temp_project_no_notes):
    """notes_update creates notes.md if it does not exist."""
    path = str(temp_project_no_notes["project_path"])
    result = js_api_no_notes.notes_update(path, "brand new notes")
    assert result["ok"] is True

    notes_file = temp_project_no_notes["project_path"] / "notes.md"
    assert notes_file.exists()
    assert notes_file.read_text() == "brand new notes"


def test_notes_update_round_trip(js_api, temp_project):
    """notes_update followed by notes_get returns same content."""
    path = str(temp_project["project_path"])
    js_api.notes_update(path, "round trip test")
    result = js_api.notes_get(path)
    assert result["ok"] is True
    assert result["data"] == "round trip test"


def test_notes_update_failure_preserves_existing_notes_file(
    js_api, temp_project, monkeypatch
):
    """Failed notes update leaves existing notes.md unchanged and removes temp file."""
    project_path = temp_project["project_path"]
    notes_file = project_path / "notes.md"
    notes_file.write_text("original notes", encoding="utf-8")
    original_write_text = Path.write_text

    def fail_temp_notes_write(self: Path, data: str, *args, **kwargs):
        if self.name == ".notes.md.tmp":
            raise OSError("simulated temp write failure")
        return original_write_text(self, data, *args, **kwargs)

    monkeypatch.setattr(Path, "write_text", fail_temp_notes_write)

    result = js_api.notes_update(str(project_path), "new notes")

    assert result["ok"] is False
    assert notes_file.read_text(encoding="utf-8") == "original notes"
    assert not (project_path / ".notes.md.tmp").exists()


# ── unrelated mutations remain unavailable ────────────────────────────────


def test_folder_move_still_deferred(js_api):
    """Folder moves remain SERVICE_UNAVAILABLE."""
    result = js_api.folder_move_to_prod_ready("/tmp/x")
    assert result["ok"] is False


def test_file_open_dev_skipped_off_windows(js_api):
    """File open is wired: dev-skipped (ok=true) off-Windows, no os.startfile (Req 6.5)."""
    result = js_api.file_open("/tmp/x")
    assert result["ok"] is True


# ── RTE format capabilities ────────────────────────────────────────────────


def test_rte_markdown_capability_and_save(js_api, temp_project):
    notes_path = temp_project["project_path"] / "notes.md"

    loaded = js_api.get_rte_file(str(notes_path))
    assert loaded["ok"] is True
    assert loaded["data"]["format"] == "markdown"
    assert loaded["data"]["capability"] == "editable"
    assert loaded["data"]["saveStrategy"] == "markdown"

    saved = js_api.save_rte_file(str(notes_path), "# Updated")
    assert saved["ok"] is True
    assert notes_path.read_text(encoding="utf-8") == "# Updated"


def test_rte_text_capability_and_plain_text_save(js_api, temp_project):
    text_path = temp_project["project_path"] / "plain.txt"
    text_path.write_text("plain", encoding="utf-8")

    loaded = js_api.get_rte_file(str(text_path))
    assert loaded["ok"] is True
    assert loaded["data"]["format"] == "text"
    assert loaded["data"]["capability"] == "editable"
    assert loaded["data"]["saveStrategy"] == "plain_text"
    assert loaded["data"]["supportedEditorFeatures"] == ["plain_text"]

    saved = js_api.save_rte_file(str(text_path), "new plain")
    assert saved["ok"] is True
    assert text_path.read_text(encoding="utf-8") == "new plain"


def test_rte_docx_is_read_only_until_source_sync_adapter_exists(js_api, temp_project):
    docx_path = temp_project["project_path"] / "_cr-docs" / "uat-signoff.docx"

    loaded = js_api.get_rte_file(str(docx_path))
    assert loaded["ok"] is True
    assert loaded["data"]["format"] == "docx"
    assert loaded["data"]["capability"] == "read_only"
    assert loaded["data"]["editable"] is False
    assert loaded["data"]["saveStrategy"] == "none"

    saved = js_api.save_rte_file(str(docx_path), "<p>must not save</p>")
    assert saved["ok"] is False
    assert saved["error"]["code"] == "RTE_SAVE_FAILED"


def test_rte_msg_is_unsupported_and_never_saved(js_api, temp_project):
    msg_path = temp_project["project_path"] / "_cr-docs" / "mail.msg"
    msg_path.parent.mkdir(parents=True, exist_ok=True)
    msg_path.write_bytes(b"not text")

    loaded = js_api.get_rte_file(str(msg_path))
    assert loaded["ok"] is True
    assert loaded["data"]["format"] == "msg"
    assert loaded["data"]["capability"] == "unsupported"
    assert loaded["data"]["content"] == ""

    saved = js_api.save_rte_file(str(msg_path), "nope")
    assert saved["ok"] is False
    assert saved["error"]["code"] == "RTE_SAVE_FAILED"


# ── RTE docx pipeline (D-0012: source.json + background export) ─────────────

PNG_1PX = (
    b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01"
    b"\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDAT\x08\xd7c\xf8\xcf"
    b"\xc0\x00\x00\x03\x01\x01\x00\x18\xdd\x8d\xb0\x00\x00\x00\x00IEND\xaeB`\x82"
)


def _doc(text: str = "hello") -> dict:
    return {
        "type": "doc",
        "content": [
            {"type": "paragraph", "content": [{"type": "text", "text": text}]}
        ],
    }


@pytest.fixture
def rte_docx(temp_project):
    """A pipeline service with a recording fake exporter + a docx target path."""
    from services.rte_document_service import RteDocumentService

    calls: list[int] = []

    def fake_exporter(source, assets_dir, target):
        calls.append(int(source.get("revision", 0)))
        target.parent.mkdir(parents=True, exist_ok=True)
        target.write_bytes(b"fake docx " + str(source.get("revision")).encode())

    service = RteDocumentService(exporter=fake_exporter)
    docx_path = temp_project["project_path"] / "_cr-docs" / "uat-signoff.docx"
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    yield {"service": service, "docx": docx_path, "calls": calls}
    service.coordinator.shutdown(timeout_s=5.0)


def test_rte_pipeline_open_empty_docx_gives_empty_doc(rte_docx):
    svc, docx = rte_docx["service"], rte_docx["docx"]
    docx.write_bytes(b"")
    opened = svc.open_document(docx)
    assert opened["needs_migration"] is False
    assert opened["revision"] == 0
    assert opened["content"] == {"type": "doc", "content": []}
    assert opened["saveStrategy"] == "docx_pipeline"
    assert opened["editable"] is True


def test_rte_pipeline_open_nonempty_docx_requests_migration(rte_docx):
    svc, docx = rte_docx["service"], rte_docx["docx"]
    from docx import Document

    d = Document()
    d.add_paragraph("legacy word content")
    d.save(str(docx))

    opened = svc.open_document(docx)
    assert opened["needs_migration"] is True
    assert opened["content"] is None
    assert "legacy word content" in (opened["content_html"] or "")


def test_rte_pipeline_stale_docx_migration_replaces_old_source(rte_docx):
    import json as jsonlib

    from docx import Document

    svc, docx = rte_docx["service"], rte_docx["docx"]
    d = Document()
    d.add_paragraph("external word edit")
    d.save(str(docx))

    old_source = svc._new_source(docx)
    old_source["revision"] = 27
    old_source["content"] = _doc("old source")
    old_source["content_hash"] = "old-hash"
    old_source["export"].update(
        {
            "last_exported_revision": 27,
            "last_exported_hash": "old-hash",
            "docx_mtime_ns": 1,
            "docx_size": 1,
        }
    )
    svc._store_source(docx, old_source)

    opened = svc.open_document(docx)
    assert opened["needs_migration"] is True
    assert opened["revision"] == 0
    backup = svc.source_path(docx).with_suffix(".json.bak")
    assert jsonlib.loads(backup.read_text(encoding="utf-8"))["revision"] == 27

    saved = svc.save_document(
        docx,
        {"content": _doc("external word edit"), "base_revision": opened["revision"], "reason": "migration"},
    )

    assert saved["revision"] == 1
    stored = jsonlib.loads(svc.source_path(docx).read_text(encoding="utf-8"))
    assert stored["revision"] == 1
    assert stored["content"]["content"][0]["content"][0]["text"] == "external word edit"


def test_rte_pipeline_save_revisions_hash_skip_and_stale(rte_docx):
    from services.rte_document_service import StaleRevisionError

    svc, docx = rte_docx["service"], rte_docx["docx"]
    first = svc.save_document(docx, {"content": _doc("a"), "base_revision": 0})
    assert first["revision"] == 1 and first["skipped"] is False
    assert svc.source_path(docx).is_file()

    same = svc.save_document(
        docx, {"content": _doc("a"), "base_revision": 1, "reason": "autosave"}
    )
    assert same["skipped"] is True and same["revision"] == 1

    with pytest.raises(StaleRevisionError):
        svc.save_document(docx, {"content": _doc("b"), "base_revision": 0})


def test_rte_pipeline_hides_sidecar_dir_on_windows(rte_docx):
    import ctypes
    import sys

    if sys.platform != "win32":
        pytest.skip("Windows hidden attribute only")

    svc, docx = rte_docx["service"], rte_docx["docx"]
    svc.save_document(docx, {"content": _doc("hidden"), "base_revision": 0})
    sidecar = svc.sidecar_dir(docx)

    attrs = ctypes.windll.kernel32.GetFileAttributesW(str(sidecar))
    assert attrs != -1
    assert attrs & 0x02


def test_rte_pipeline_save_rejected_for_implemented_project(tmp_path):
    from services.rte_document_service import RteDocumentService

    svc = RteDocumentService(exporter=lambda *a: None)
    docx = tmp_path / "2024" / "IMPLEMENTED" / "proj" / "_cr-docs" / "uat-signoff.docx"
    docx.parent.mkdir(parents=True)
    opened = svc.open_document(docx)
    assert opened["capability"] == "read_only"
    assert opened["saveStrategy"] == "none"
    with pytest.raises(ValueError):
        svc.save_document(docx, {"content": _doc(), "base_revision": 0})
    svc.coordinator.shutdown(timeout_s=2.0)


def test_rte_pipeline_image_validation_and_traversal_guard(rte_docx):
    import base64 as b64

    svc, docx = rte_docx["service"], rte_docx["docx"]
    saved = svc.save_image(docx, b64.b64encode(PNG_1PX).decode())
    assert saved["src"].startswith("asset://")
    assert (svc.assets_dir(docx) / saved["file_name"]).is_file()

    # duplicate paste dedupes to the same content-addressed file
    again = svc.save_image(docx, b64.b64encode(PNG_1PX).decode())
    assert again["asset_id"] == saved["asset_id"]

    with pytest.raises(ValueError):
        svc.save_image(docx, b64.b64encode(b"plain text bytes").decode())
    with pytest.raises(ValueError):
        svc.read_asset(docx, "asset://../../secret.png")
    with pytest.raises(ValueError):
        svc.read_asset(docx, "..\\..\\secret.png")

    read = svc.read_asset(docx, saved["src"])
    assert read["data_uri"].startswith("data:image/png;base64,")


def test_rte_pipeline_dehydration_sweeps_loose_base64_images(rte_docx):
    import base64 as b64
    import json as jsonlib

    svc, docx = rte_docx["service"], rte_docx["docx"]
    data_uri = "data:image/png;base64," + b64.b64encode(PNG_1PX).decode()
    content = {
        "type": "doc",
        "content": [
            {
                "type": "paragraph",
                "content": [{"type": "image", "attrs": {"src": data_uri}}],
            }
        ],
    }
    svc.save_document(docx, {"content": content, "base_revision": 0})
    stored = jsonlib.loads(svc.source_path(docx).read_text(encoding="utf-8"))
    img = stored["content"]["content"][0]["content"][0]
    assert img["attrs"]["src"].startswith("asset://")
    assert img["attrs"]["assetId"]
    assert list(svc.assets_dir(docx).glob("*.png"))


def test_rte_pipeline_coordinator_latest_revision_wins(temp_project):
    import threading as th

    from services.rte_document_service import RteDocumentService

    started = th.Event()
    release = th.Event()
    calls: list[int] = []

    def slow_exporter(source, assets_dir, target):
        calls.append(int(source.get("revision", 0)))
        started.set()
        release.wait(timeout=10)

    svc = RteDocumentService(exporter=slow_exporter)
    docx = temp_project["project_path"] / "_cr-docs" / "prod-lv.docx"
    docx.parent.mkdir(parents=True, exist_ok=True)

    svc.save_document(docx, {"content": _doc("r1"), "base_revision": 0, "reason": "manual"})
    assert started.wait(timeout=5)
    # While rev-1 export runs, revisions 2..4 arrive: only 4 must export.
    svc.save_document(docx, {"content": _doc("r2"), "base_revision": 1, "reason": "manual"})
    svc.save_document(docx, {"content": _doc("r3"), "base_revision": 2, "reason": "manual"})
    svc.save_document(docx, {"content": _doc("r4"), "base_revision": 3, "reason": "manual"})
    release.set()
    assert svc.coordinator.wait_idle(timeout_s=10)
    svc.coordinator.shutdown(timeout_s=5.0)
    assert calls[0] == 1
    assert calls[-1] == 4
    assert len(calls) == 2, f"intermediate revisions must be dropped, got {calls}"


def test_rte_pipeline_locked_target_marks_pending_and_keeps_old_docx(temp_project):
    from infrastructure.docx_writer import DocxTargetLockedError
    from services.rte_document_service import RteDocumentService

    def locked_exporter(source, assets_dir, target):
        raise DocxTargetLockedError(str(target))

    svc = RteDocumentService(exporter=locked_exporter)
    docx = temp_project["project_path"] / "_cr-docs" / "uat-signoff.docx"
    docx.parent.mkdir(parents=True, exist_ok=True)
    docx.write_bytes(b"previous word file")

    svc.save_document(docx, {"content": _doc("x"), "base_revision": 0, "reason": "manual"})
    assert svc.coordinator.wait_idle(timeout_s=10)
    status = svc.export_status(docx)
    assert status["state"] == "pending_retry"
    assert status["export_pending"] is True
    assert docx.read_bytes() == b"previous word file"
    svc.coordinator.shutdown(timeout_s=5.0)


def test_rte_pipeline_real_export_roundtrip(temp_project):
    """End-to-end with the real python-docx exporter."""
    from docx import Document

    from services.rte_document_service import RteDocumentService

    svc = RteDocumentService()
    docx = temp_project["project_path"] / "_cr-docs" / "uat-signoff.docx"
    docx.parent.mkdir(parents=True, exist_ok=True)

    svc.save_document(
        docx, {"content": _doc("exported body"), "base_revision": 0, "reason": "manual"}
    )
    assert svc.coordinator.wait_idle(timeout_s=15)
    svc.coordinator.shutdown(timeout_s=5.0)
    assert docx.is_file() and docx.stat().st_size > 0
    texts = [p.text for p in Document(str(docx)).paragraphs]
    assert any("exported body" in t for t in texts)
